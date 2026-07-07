#!/usr/bin/env python3
"""docx-platform-tests adapter protocol v1 entrypoint for python-docx.

Capability map (see README.md):
- replaceFirstTextOccurrence: intra-run replacement via the library's
  paragraph/run API. A match that spans runs is declined (exit 2) rather
  than approximated with adapter-side algorithms.
- acceptAllTrackedChanges / rejectAllTrackedChanges and everything else:
  exit 2 unsupported -- python-docx has no revision API.
"""
import argparse
import json
import sys

PROTOCOL_VERSION = "1"


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--protocol-version", required=True)
    parser.add_argument("--operation", required=True)
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    if args.protocol_version != PROTOCOL_VERSION:
        print(
            f"python-docx adapter speaks protocol v{PROTOCOL_VERSION}, "
            f"got {args.protocol_version}"
        )
        return 3

    with open(args.operation, encoding="utf-8") as f:
        operation = json.load(f)
    name = operation.get("operationName")

    import docx  # imported late so unsupported paths need no library
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph

    def apply_run_formatting(run, formatting):
        if not formatting:
            return
        if "bold" in formatting:
            run.bold = bool(formatting["bold"])
        if "italic" in formatting:
            run.italic = bool(formatting["italic"])
        if "fontSizeHalfPoints" in formatting:
            run.font.size = Pt(formatting["fontSizeHalfPoints"] / 2)

    def paragraph_containing(document, text):
        for paragraph in document.paragraphs:
            if text in paragraph.text:
                return paragraph
        return None

    def apply_numbering(paragraph, num_id, ilvl):
        p_pr = paragraph._p.get_or_add_pPr()
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            num_pr = OxmlElement("w:numPr")
            p_pr.append(num_pr)
        ilvl_el = num_pr.find(qn("w:ilvl"))
        if ilvl_el is None:
            ilvl_el = OxmlElement("w:ilvl")
            num_pr.append(ilvl_el)
        ilvl_el.set(qn("w:val"), str(ilvl))
        num_id_el = num_pr.find(qn("w:numId"))
        if num_id_el is None:
            num_id_el = OxmlElement("w:numId")
            num_pr.append(num_id_el)
        num_id_el.set(qn("w:val"), str(num_id))

    def replace_in_runs(document, find_text, replace_text, formatting=None):
        for paragraph in document.paragraphs:
            if find_text not in paragraph.text:
                continue
            for run in paragraph.runs:
                if find_text in run.text:
                    run.text = run.text.replace(find_text, replace_text, 1)
                    apply_run_formatting(run, formatting)
                    return 0
            print(
                "match spans run boundaries; the python-docx adapter only "
                "performs intra-run replacement/formatting (glue, not algorithms)"
            )
            return 2
        print(f"findText not present in any paragraph: {find_text!r}", file=sys.stderr)
        return 1

    if name in ("acceptAllTrackedChanges", "rejectAllTrackedChanges"):
        print("python-docx has no tracked-changes (revision) API")
        return 2

    if name == "composeDocumentWithParagraphs":
        document = docx.Document()
        for descriptor in operation["paragraphDescriptorList"]:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(descriptor["paragraphText"])
            apply_run_formatting(run, descriptor.get("runFormatting"))
        document.save(args.output)
        return 0

    if name == "composeDocumentWithTable":
        rows = operation["tableCellTextRows"]
        document = docx.Document()
        table = document.add_table(rows=len(rows), cols=len(rows[0]))
        for r_idx, row in enumerate(rows):
            for c_idx, text in enumerate(row):
                table.cell(r_idx, c_idx).text = text
        document.save(args.output)
        return 0

    if name == "composeDocumentWithHeaderText":
        document = docx.Document()
        document.add_paragraph(operation["bodyText"])
        section = document.sections[0]
        section.header.paragraphs[0].text = operation["headerText"]
        document.save(args.output)
        return 0

    if name == "composeDocumentWithNumberedList":
        document = docx.Document()
        if operation.get("numberFormat") != "decimal":
            print("python-docx adapter only maps composeDocumentWithNumberedList decimal to List Number style")
            return 2
        for text in operation["listItemTexts"]:
            document.add_paragraph(text, style="List Number")
        document.save(args.output)
        return 0

    if name == "composeDocumentWithHyperlink":
        print("python-docx exposes no public API for writing external hyperlinks")
        return 2

    document = docx.Document(args.input)

    if name == "replaceFirstTextOccurrence":
        result = replace_in_runs(document, operation["findText"], operation["replaceText"])
        if result == 0:
            document.save(args.output)
        return result

    if name == "appendParagraphWithText":
        paragraph = document.add_paragraph()
        run = paragraph.add_run(operation["paragraphText"])
        apply_run_formatting(run, operation.get("runFormatting"))
        document.save(args.output)
        return 0

    if name == "insertParagraphAfterAnchorText":
        paragraph = paragraph_containing(document, operation["anchorText"])
        if paragraph is None:
            print(f"anchorText not present in any paragraph: {operation['anchorText']!r}", file=sys.stderr)
            return 1
        new_paragraph = OxmlElement("w:p")
        paragraph._p.addnext(new_paragraph)
        inserted = Paragraph(new_paragraph, paragraph._parent)
        inserted.add_run(operation["paragraphText"])
        document.save(args.output)
        return 0

    if name == "appendTableRow":
        table = document.tables[int(operation["tableIndex"])]
        row = table.add_row()
        for idx, text in enumerate(operation["cellTexts"]):
            row.cells[idx].text = text
        document.save(args.output)
        return 0

    if name == "deleteTableRowAtIndex":
        print("python-docx has no public table row deletion API")
        return 2

    if name == "setTableCellText":
        cell = document.tables[int(operation["tableIndex"])].cell(
            int(operation["rowIndex"]), int(operation["columnIndex"])
        )
        cell.text = operation["replacementCellText"]
        document.save(args.output)
        return 0

    if name == "mergeTableCellsInRow":
        table = document.tables[int(operation["tableIndex"])]
        row_idx = int(operation["rowIndex"])
        start = int(operation["startColumnIndex"])
        end = int(operation["endColumnIndex"])
        table.cell(row_idx, start).merge(table.cell(row_idx, end))
        document.save(args.output)
        return 0

    if name == "addCommentOnFirstTextOccurrence":
        paragraph = paragraph_containing(document, operation["anchorText"])
        if paragraph is None:
            print(f"anchorText not present in any paragraph: {operation['anchorText']!r}", file=sys.stderr)
            return 1
        runs = [run for run in paragraph.runs if operation["anchorText"] in run.text]
        if not runs:
            print("comment anchor spans run boundaries; python-docx add_comment requires run anchors")
            return 2
        if not hasattr(document, "add_comment"):
            print("installed python-docx exposes no add_comment API")
            return 2
        document.add_comment(
            runs,
            text=operation["commentText"],
            author=operation["commentAuthorName"],
            initials=operation["commentAuthorInitials"],
        )
        document.save(args.output)
        return 0

    if name == "removeAllComments":
        print("python-docx exposes no public API to remove all comments and anchors")
        return 2

    if name == "applyParagraphStyleToAnchor":
        paragraph = paragraph_containing(document, operation["anchorText"])
        if paragraph is None:
            print(f"anchorText not present in any paragraph: {operation['anchorText']!r}", file=sys.stderr)
            return 1
        paragraph.style = operation["paragraphStyleId"]
        document.save(args.output)
        return 0

    if name == "applyNumberingToAnchorParagraph":
        paragraph = paragraph_containing(document, operation["anchorText"])
        if paragraph is None:
            print(f"anchorText not present in any paragraph: {operation['anchorText']!r}", file=sys.stderr)
            return 1
        apply_numbering(paragraph, operation["numberingInstanceId"], operation["indentationLevel"])
        document.save(args.output)
        return 0

    if name == "formatFirstTextOccurrence":
        result = replace_in_runs(
            document,
            operation["findText"],
            operation["findText"],
            operation.get("runFormatting"),
        )
        if result == 0:
            document.save(args.output)
        return result

    if name == "setDefaultFooterText":
        document.sections[0].footer.paragraphs[0].text = operation["footerText"]
        document.save(args.output)
        return 0

    print(f"python-docx adapter does not implement operation '{name}'")
    return 2


if __name__ == "__main__":
    sys.exit(main())
